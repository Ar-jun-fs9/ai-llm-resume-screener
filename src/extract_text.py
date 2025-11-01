import os
from pdfminer.high_level import extract_text as extract_pdf_text
import docx
from io import BytesIO
import tempfile


def extract_text_from_pdf(pdf_file):
    try:
        # pdfminer expects a file path or file-like object
        if hasattr(pdf_file, "read"):
            # It's a file-like object (Flask FileStorage or similar)
            pdf_bytes = pdf_file.read()
            pdf_io = BytesIO(pdf_bytes)
            text = extract_pdf_text(pdf_io)
        else:
            # Assume it's a path
            text = extract_pdf_text(pdf_file)
        return text
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""


def extract_text_from_docx(docx_file):
    try:
        if hasattr(docx_file, "read"):
            # It's a file-like object (Flask FileStorage or similar)
            # Save to temporary file since python-docx needs a file path
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(docx_file.read())
                temp_path = temp_file.name

            doc = docx.Document(temp_path)
            # Clean up temporary file
            os.unlink(temp_path)
        else:
            # Assume it's a path
            doc = docx.Document(docx_file)

        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return "\n".join(fullText)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""


def extract_text(file):
    if hasattr(file, "filename"):
        # Flask FileStorage has filename attribute
        filename = file.filename
        if filename.endswith(".pdf"):
            return extract_text_from_pdf(file)
        elif filename.endswith(".docx"):
            return extract_text_from_docx(file)
        else:
            print(f"Unsupported file format: {filename}")
            return ""
    elif hasattr(file, "name"):
        # Other file-like objects (including Streamlit UploadedFile)
        filename = file.name
        if filename.endswith(".pdf"):
            return extract_text_from_pdf(file)
        elif filename.endswith(".docx"):
            return extract_text_from_docx(file)
        else:
            print(f"Unsupported file format: {filename}")
            return ""
    else:
        # Fallback for file paths
        if file.endswith(".pdf"):
            return extract_text_from_pdf(file)
        elif file.endswith(".docx"):
            return extract_text_from_docx(file)
        else:
            print(f"Unsupported file format: {file}")
            return ""
